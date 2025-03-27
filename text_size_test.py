from pypdf import PdfReader
from docx import Document
from tokenCounter import count_tokens
from modelinfo import get_context_length
import re
import ollama
import argparse
import pypdf.errors
import pytest

def extract(filepath:str) -> str:
    """
    Extract entire pdf text
    """
    try:
        reader = PdfReader(filepath)
        pages = reader.pages

        extracted_text = ""
        
        num_pages = 0 # testing number of pages processed
        for i in range(len(pages)):
            curr_text = pages[i].extract_text()

            if not curr_text:
                print("No text found on this page.") 
                continue
            
            print(f"PROCESSING PAGE {i+1}\n\n")
            num_pages += 1
            curr_text = curr_text.encode('utf-8', 'replace').decode('utf-8')

            extracted_text += f"PAGE {i+1}\n\n" if i + 1 == 1 else f"\n\nPAGE {i+1}\n\n"
            extracted_text += curr_text

        return extracted_text
    except FileNotFoundError:
        print("Error: PDF file not found.")
        return ""
    except pypdf.errors.PdfReadError as e:
        print(f"Error while reading file: {e}")
        return ""
    except ValueError as e:
        print(f"ValueError: {e}")
    except Exception as e:
        print(f"Unexpected error: {e}")

def chunk_text(text:str, num_tokens:int, overlap: float = 0.3) -> list[str]:
    """
    Splits provided text by the number of tokens with overlapping regions.
    """
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks = []
    curr_chunk = []
    curr_length = 0

    for sentence in sentences:
        sentence_length = count_tokens(text=sentence)

        if sentence_length + curr_length <= num_tokens:
            curr_chunk.append(sentence)
            curr_length += sentence_length
        else:
            chunks.append(" ".join(curr_chunk))

            overlap_size = int(overlap * num_tokens)

            retained_tokens = []
            retained_length = 0
            while curr_chunk and retained_length < overlap_size:
                retained_sentence = curr_chunk.pop()
                retained_tokens.insert(0, retained_sentence)
                retained_length += count_tokens(text=retained_sentence)

            curr_chunk = retained_tokens + [sentence]
            curr_length = retained_length + sentence_length

    if curr_chunk:
        chunks.append(" ".join(curr_chunk))

    return chunks

#TODO: handle NoneType model response and fix
#TODO: fix memory usage issue (VRAM) - mistral
def get_summary(text:str, prompt:str, model:str) -> str:
    """
    Prompt the model based on the provided chunk of a text.
    """
    try:
        model_response = ollama.chat(
            model=model,
            messages=[
                {'role':'user', 'content':f'{prompt}\n{text}'}
            ],
            stream=False,
            options={'temperature':0, 'num_ctx': get_context_length(model)}
        )
        
        return model_response.get('message', {}).get('content', "Error: no response from the model.")
    except Exception as e:
        return f"Error: {str(e)}"

#TODO: Find a way to automate test recoding into docx (maybe LaTeX?)
#TODO: Figure out a way to parse chunk texts and model's responses as well
def record_test(text:str, prompt:str, chunk_prompt:str, model:str, response:str) -> None:
    """
    Records test in a txt file with text, prompt, model, and model's response
    """

    test_dir = "tests/text_size_test/"

    model_prefix = {
        "llama3.1:8b": "llama/",
        "llama3.1:8b-instruct-fp16": "llama/",
        "qwen2.5:14b": "qwen/",
        "qwen2.5:14b-instruct-fp16": "qwen/",
        "mistral-nemo:12b": "mistral/",
        "mistral-nemo:12b-instruct-2407-fp16": "mistral/",
        "gemma3:27b":"gemma/",
        "phi4:14b":"phi4/",
        "deepseek-r1:14b":"deepseek/",
        "deepseek-r1:14b-qwen-distill-fp16":"deepseek/",
    }.get(model, "unknown/")

    test_dir += model_prefix
    is_instruct = "instruct" in model
    model_type = "instruct" if is_instruct else "base"
    result_file = f"{test_dir}{model_type}_modelresponse.docx"

    # SAVING TO DOCX
    document = Document()
    
    document.add_heading(f"{model} Test", 0)

    document.add_heading(f"Параметры модели:\n", level=3)
    document.add_paragraph("Температура: 0", style='List Bullet')
    document.add_paragraph(f"num_ctx: {get_context_length(model)}\n", style='List Bullet')
    
    document.add_heading("Подход\n", level=2)
    document.add_paragraph("<Описание подхода с параметрами для разбития текста (ко-во токнов, оверлап и тд)>")
    document.add_paragraph(f"Общий размер финального текста вместе с промптом: {count_tokens(text=text) + count_tokens(text=prompt)}")

    document.add_heading("Промпты\n", level=2)
    document.add_paragraph(f"Общий промпт:\n", style="List Bullet")
    document.add_paragraph(f"{prompt}\n")
    document.add_paragraph(f"Промпт для чанков: \n", style="List Bullet")
    document.add_paragraph(f"{chunk_prompt}\n")

    document.add_heading("Финальный Ответ Модели\n", level=2)
    document.add_paragraph(f"Ответ модели: \n\n")
    document.add_paragraph(f"{response}\n\n")
    document.add_paragraph(f"Финальный текст:\n\n ")
    document.add_paragraph(f"{text}\n")
    

    document.add_page_break()

    document.add_heading("Сравнение с Claude", level=0)
    document.save(result_file)


# Parsing args and recording tests
if __name__ == "__main__":
    chunk_prompt = """Conduct a holistic, precision-driven text comprehension analysis. Your goal is to maintain 100% information integrity while preserving exact contextual nuances.

Deliver:

1. An exhaustive factual summary capturing all key events and details without distortion or omission.
2. A thematic analysis exploring core ideas and recurring motifs.
3.A deep symbolic/metaphorical interpretation supported by textual evidence.
4. Extraction of five pivotal narrative components with justification.

Map out with full detail:

1. All character interactions and relationships.
2. The complete plot progression with no missing elements.
3. Narrative techniques, including structural choices, perspective, and literary devices.
4. A linguistic breakdown covering syntax, diction, tone, narrative voice, and literary techniques.
"""
    final_prompt = """Conduct a holistic, precision-driven analysis of the following summaries. Your goal is to maintain 100% information integrity while preserving exact contextual nuances.

Deliver:

1. An exhaustive factual summary capturing all key events and details without distortion or omission.
2. A thematic analysis exploring core ideas and recurring motifs.
3. A deep symbolic/metaphorical interpretation supported by textual evidence.
4. Extraction of five pivotal narrative components with justification.

Map out with full detail:

1. All character interactions and relationships.
2. The complete plot progression with no missing elements.
3. Narrative techniques, including structural choices, perspective, and literary devices.
4. A linguistic breakdown covering syntax, diction, tone, narrative voice, and literary techniques.
"""

    parser = argparse.ArgumentParser(description="User input")
    parser.add_argument("num_tokens", type=int, help="Number of tokens to process at a time")
    parser.add_argument("filepath", type=str, help="Filepath to pdf")
    args = parser.parse_args()

    model_name = "phi4:14b"

    print(f"Processing model with {args.num_tokens} tokens: {model_name}\n\n")

    extracted_text = extract(args.filepath)
    chunks = chunk_text(extracted_text, args.num_tokens)

    chunk_summaries = []
    for chunk in chunks:
        summary = get_summary(chunk, chunk_prompt, model_name)
        chunk_summaries.append(summary)

    final_summary = get_summary(" ".join(chunk_summaries), final_prompt, model_name)

    record_test(" ".join(chunk_summaries), final_prompt, chunk_prompt, model_name, final_summary)

    # response: ollama.ListResponse = ollama.list()
    # for model in response.models:
    #     model_name = model.model

    #     if model_name == "nomic-embed-text:latest":
    #         continue

    #     print(f"Processing model with {args.num_tokens} tokens: {model_name}\n")
        
    #     extracted_text = extract(args.filepath)

    #     chunks = chunk_text(extracted_text, args.num_tokens)

    #     chunk_summaries = []
    #     for chunk in chunks:
    #         summary = get_summary(chunk, chunk_prompt, model_name)
    #         chunk_summaries.append(summary)
        
    #     final_summary = get_summary(" ".join(chunk_summaries), final_prompt, model_name)

    #     record_test(" ".join(chunk_summaries), final_prompt, model_name, final_summary)








# def test_extract():
#     #TESTING extract()
#     prompt = """Conduct a holistic, precision-driven text comprehension analysis. Your goal is to maintain 100% information integrity while preserving exact contextual nuances.

# Deliver:

# 1. An exhaustive factual summary capturing all key events and details without distortion or omission.
# 2. A thematic analysis exploring core ideas and recurring motifs.
# 3.A deep symbolic/metaphorical interpretation supported by textual evidence.
# 4. Extraction of five pivotal narrative components with justification.

# Map out with full detail:

# 1. All character interactions and relationships.
# 2. The complete plot progression with no missing elements.
# 3. Narrative techniques, including structural choices, perspective, and literary devices.
# 4. A linguistic breakdown covering syntax, diction, tone, narrative voice, and literary techniques.
# """
#     filepath = "tests/text_size_test/Rye.pdf"
#     num_tokens = 8192
#     extracted_text, num_pages = extract(filepath, num_tokens, prompt)

#     with open("tests/text_size_test/extract_test.txt", 'w', encoding="utf-8") as file:
#         file.write(extracted_text)
    
#     assert num_pages >= 10, f"Expected at least 10 pages, but got {num_pages}"
#     assert len(extracted_text) > 0, "Extracted text should not be empty"